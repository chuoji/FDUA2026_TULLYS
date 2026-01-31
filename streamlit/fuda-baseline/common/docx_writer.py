"""
.docx生成モジュール
提案書と検証レポートを.docx/.txt形式で生成
"""

import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime


def create_proposal_docx(company_code: int, proposal_text: str) -> bytes:
    """
    提案書.docx生成

    Args:
        company_code: 企業コード
        proposal_text: 提案書本文（Markdown形式）

    Returns:
        bytes: .docxファイルのバイナリデータ
    """
    doc = Document()

    # タイトル
    title = doc.add_heading(f'成長戦略提案書', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # サブタイトル
    subtitle = doc.add_paragraph(f'企業コード: {company_code}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.size = Pt(14)

    # 日付
    date_para = doc.add_paragraph(
        f'作成日: {datetime.datetime.now().strftime("%Y年%m月%d日")}'
    )
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 改ページ
    doc.add_page_break()

    # 本文を段落に分割して追加
    lines = proposal_text.split('\n')

    for line in lines:
        line = line.strip()

        if not line:
            # 空行
            doc.add_paragraph()
            continue

        # Markdownヘッダーを検出
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('- ') or line.startswith('* '):
            # 箇条書き
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith(tuple(f'{i}. ' for i in range(1, 10))):
            # 番号付きリスト
            # 数字の後のドットとスペースを除去
            content = line.split('. ', 1)[1] if '. ' in line else line
            doc.add_paragraph(content, style='List Number')
        else:
            # 通常の段落
            para = doc.add_paragraph(line)
            para_format = para.paragraph_format
            para_format.line_spacing = 1.5

    # バイナリデータとして返す
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def create_verification_report_txt(company_code: int, validation_result: dict) -> str:
    """
    検証レポート.txt生成

    Args:
        company_code: 企業コード
        validation_result: 検証結果の辞書

    Returns:
        str: 検証レポートのテキスト
    """
    lines = []
    lines.append("=" * 80)
    lines.append(f"検証プロセスレポート - 企業コード {company_code}")
    lines.append(f"作成日時: {datetime.datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
    lines.append("=" * 80)
    lines.append("")

    # 1. 形式チェック
    lines.append("## 1. 形式チェック")
    lines.append("-" * 80)
    format_check = validation_result.get('format', {})
    lines.append(f"文字数: {format_check.get('char_count', 0):,}字 "
                 f"(目標: 15,000字以内)")
    lines.append(f"必須セクション: {format_check.get('has_required_sections', False)}")
    if format_check.get('missing_sections'):
        lines.append(f"  欠落セクション: {', '.join(format_check['missing_sections'])}")
    lines.append(f"判定: {'✓ 合格' if format_check.get('passed', False) else '✗ 不合格'}")
    lines.append("")

    # 2. 数値整合性チェック
    lines.append("## 2. 数値整合性チェック")
    lines.append("-" * 80)
    numerical_check = validation_result.get('numerical', {})
    lines.append(f"検出された数値: {numerical_check.get('numbers_found', 0)}個")
    lines.append(f"範囲外の数値: {numerical_check.get('out_of_range', 0)}個")
    if numerical_check.get('issues'):
        lines.append("  問題箇所:")
        for issue in numerical_check['issues']:
            lines.append(f"    - {issue}")
    lines.append(f"判定: {'✓ 合格' if numerical_check.get('passed', False) else '✗ 不合格'}")
    lines.append("")

    # 3. カバレッジチェック
    lines.append("## 3. カバレッジチェック")
    lines.append("-" * 80)
    coverage_check = validation_result.get('coverage', {})
    lines.append(f"財務データ引用: {coverage_check.get('financial_refs', 0)}件")
    lines.append(f"PDF引用: {coverage_check.get('pdf_refs', 0)}件")
    lines.append(f"根拠の明示: {coverage_check.get('has_evidence', False)}")
    lines.append(f"判定: {'✓ 合格' if coverage_check.get('passed', False) else '✗ 不合格'}")
    lines.append("")

    # 4. 総合評価
    lines.append("## 4. 総合評価")
    lines.append("-" * 80)
    overall = validation_result.get('overall', {})
    lines.append(f"総合判定: {'✓ 合格' if overall.get('passed', False) else '✗ 不合格'}")
    lines.append(f"品質スコア: {overall.get('score', 0):.1f}/100")
    lines.append("")

    # 5. 改善提案
    if validation_result.get('suggestions'):
        lines.append("## 5. 改善提案")
        lines.append("-" * 80)
        for suggestion in validation_result['suggestions']:
            lines.append(f"- {suggestion}")
        lines.append("")

    lines.append("=" * 80)
    lines.append("レポート終了")
    lines.append("=" * 80)

    return "\n".join(lines)


def create_prompt_log_txt(logs: list[dict]) -> str:
    """
    プロンプトログ.txt生成（複数の生成ログをまとめる）

    Args:
        logs: 生成ログのリスト

    Returns:
        str: プロンプトログのテキスト
    """
    lines = []
    lines.append("=" * 80)
    lines.append("FDUA Competition - 提案書生成ログ")
    lines.append(f"出力日時: {datetime.datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
    lines.append(f"総生成数: {len(logs)}件")
    lines.append("=" * 80)
    lines.append("")

    for i, log in enumerate(logs, 1):
        lines.append(f"{'=' * 80}")
        lines.append(f"生成 #{i}")
        lines.append(f"{'=' * 80}")
        lines.append(f"タイムスタンプ: {log.get('timestamp', 'N/A')}")
        lines.append(f"企業コード: {log.get('company_code', 'N/A')}")
        lines.append(f"企業情報: {log.get('company_info', {})}")
        lines.append(f"所要時間: {log.get('duration_seconds', 0):.1f}秒")
        lines.append("")
        lines.append("[プロンプト]")
        lines.append("-" * 80)
        lines.append(log.get('prompt', ''))
        lines.append("")
        lines.append("[応答]")
        lines.append("-" * 80)
        lines.append(log.get('response', ''))
        lines.append("")
        lines.append("[検証結果]")
        lines.append("-" * 80)
        lines.append(str(log.get('validation', {})))
        lines.append("")
        lines.append("")

    return "\n".join(lines)
